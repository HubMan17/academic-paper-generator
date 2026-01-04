from dataclasses import asdict

from django.shortcuts import render
from rest_framework import status
from rest_framework.decorators import api_view
from rest_framework.response import Response
from drf_spectacular.utils import extend_schema

from services.analyzer import RepoAnalyzer
from services.llm import LLMClient
from services.llm.errors import LLMError
from services.prompting import slice_for_section
from apps.llm.models import LLMCall
from .serializers import (
    AnalyzeRequestSerializer,
    AnalyzeResponseSerializer,
    LLMTextRequestSerializer,
    LLMJsonRequestSerializer,
    LLMResponseSerializer,
)


def test_page(request):
    return render(request, 'core/test_page.html')


def test_analyzer_page(request):
    return render(request, 'core/test_analyzer.html')


def dev_pipeline_page(request):
    return render(request, 'dev/pipeline.html')


@extend_schema(
    request=AnalyzeRequestSerializer,
    responses={200: AnalyzeResponseSerializer},
    description="Анализирует репозиторий и возвращает facts.json",
    tags=["Analyzer"]
)
@api_view(['POST'])
def analyze_repository_api(request):
    serializer = AnalyzeRequestSerializer(data=request.data)
    if not serializer.is_valid():
        return Response(
            {"status": "error", "error": serializer.errors, "facts": None},
            status=status.HTTP_400_BAD_REQUEST
        )

    repo_url = serializer.validated_data['repo_url']

    try:
        analyzer = RepoAnalyzer(repo_url)
        facts = analyzer.analyze()
        return Response({
            "status": "success",
            "facts": facts,
            "error": None
        })
    except Exception as e:
        return Response(
            {"status": "error", "error": str(e), "facts": None},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@extend_schema(
    request=LLMTextRequestSerializer,
    responses={200: LLMResponseSerializer},
    description="Генерирует текст через LLM",
    tags=["LLM"]
)
@api_view(['POST'])
def llm_generate_text_api(request):
    serializer = LLMTextRequestSerializer(data=request.data)
    if not serializer.is_valid():
        return Response(
            {"status": "error", "error": serializer.errors, "result": None},
            status=status.HTTP_400_BAD_REQUEST
        )

    data = serializer.validated_data

    try:
        client = LLMClient()
        result = client.generate_text(
            system=data['system'],
            user=data['user'],
            model=data.get('model'),
            temperature=data.get('temperature', 0.7),
            use_cache=data.get('use_cache', True),
        )
        return Response({
            "status": "success",
            "result": {
                "text": result.text,
                "meta": asdict(result.meta),
            },
            "error": None
        })
    except LLMError as e:
        return Response(
            {"status": "error", "error": str(e), "result": None},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@extend_schema(
    request=LLMJsonRequestSerializer,
    responses={200: LLMResponseSerializer},
    description="Генерирует JSON через LLM",
    tags=["LLM"]
)
@api_view(['POST'])
def llm_generate_json_api(request):
    serializer = LLMJsonRequestSerializer(data=request.data)
    if not serializer.is_valid():
        return Response(
            {"status": "error", "error": serializer.errors, "result": None},
            status=status.HTTP_400_BAD_REQUEST
        )

    data = serializer.validated_data

    try:
        client = LLMClient()
        result = client.generate_json(
            system=data['system'],
            user=data['user'],
            model=data.get('model'),
            temperature=data.get('temperature', 0.3),
            use_cache=data.get('use_cache', True),
            schema=data.get('schema'),
        )
        return Response({
            "status": "success",
            "result": {
                "data": result.data,
                "meta": asdict(result.meta),
            },
            "error": None
        })
    except LLMError as e:
        return Response(
            {"status": "error", "error": str(e), "result": None},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@extend_schema(
    responses={200: dict},
    description="Получает статистику LLM вызовов",
    tags=["LLM"]
)
@api_view(['GET'])
def llm_stats_api(request):
    total = LLMCall.objects.count()
    success = LLMCall.objects.filter(status=LLMCall.Status.SUCCESS).count()
    failed = LLMCall.objects.filter(status=LLMCall.Status.FAILED).count()
    in_progress = LLMCall.objects.filter(status=LLMCall.Status.IN_PROGRESS).count()

    recent = LLMCall.objects.order_by('-created_at')[:10]
    recent_list = [
        {
            "fingerprint": call.fingerprint[:16] + "...",
            "model": call.model,
            "status": call.status,
            "created_at": call.created_at.isoformat(),
            "cost": call.meta.get("cost_estimate", 0) if call.meta else 0,
        }
        for call in recent
    ]

    return Response({
        "total": total,
        "success": success,
        "failed": failed,
        "in_progress": in_progress,
        "recent": recent_list,
    })


@extend_schema(
    responses={200: dict},
    description="Очищает кэш LLM вызовов",
    tags=["LLM"]
)
@api_view(['POST'])
def llm_clear_cache_api(request):
    deleted, _ = LLMCall.objects.all().delete()
    return Response({
        "status": "success",
        "deleted": deleted,
    })


@extend_schema(
    responses={200: dict},
    description="Создаёт ContextPack для секции документа",
    tags=["Prompting"]
)
@api_view(['POST'])
def prompting_slice_api(request):
    section_key = request.data.get('section_key')
    facts = request.data.get('facts', {})
    outline = request.data.get('outline', {})
    summaries = request.data.get('summaries', [])
    global_context = request.data.get('global_context', '')

    if not section_key:
        return Response(
            {"status": "error", "error": "section_key is required", "context_pack": None},
            status=status.HTTP_400_BAD_REQUEST
        )

    try:
        context_pack = slice_for_section(
            section_key=section_key,
            facts=facts,
            outline=outline,
            summaries=summaries,
            global_context=global_context
        )

        pack_data = {
            "section_key": context_pack.section_key,
            "layers": {
                "global_context": context_pack.layers.global_context,
                "outline_excerpt": context_pack.layers.outline_excerpt,
                "facts_slice": context_pack.layers.facts_slice,
                "summaries": context_pack.layers.summaries,
                "constraints": context_pack.layers.constraints
            },
            "rendered_prompt": {
                "system": context_pack.rendered_prompt.system,
                "user": context_pack.rendered_prompt.user
            },
            "budget": {
                "max_input_tokens_approx": context_pack.budget.max_input_tokens_approx,
                "max_output_tokens": context_pack.budget.max_output_tokens,
                "soft_char_limit": context_pack.budget.soft_char_limit
            },
            "debug": {
                "selected_fact_refs": [
                    {"fact_id": ref.fact_id, "reason": ref.reason, "weight": ref.weight}
                    for ref in context_pack.debug.selected_fact_refs
                ],
                "selection_reason": context_pack.debug.selection_reason,
                "trims_applied": context_pack.debug.trims_applied
            }
        }

        return Response({
            "status": "success",
            "context_pack": pack_data,
            "error": None
        })
    except ValueError as e:
        return Response(
            {"status": "error", "error": str(e), "context_pack": None},
            status=status.HTTP_400_BAD_REQUEST
        )
    except Exception as e:
        return Response(
            {"status": "error", "error": str(e), "context_pack": None},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


def test_pipeline_page(request):
    return render(request, 'core/test_pipeline.html')


@extend_schema(
    responses={200: dict},
    description="Запускает полный пайплайн генерации документа",
    tags=["Pipeline"]
)
@api_view(['POST'])
def pipeline_run_api(request):
    work_type = request.data.get('work_type', 'referat')
    profile = request.data.get('profile', 'default')
    topic_title = request.data.get('topic_title', 'Разработка программного обеспечения')
    topic_description = request.data.get('topic_description', '')
    facts = request.data.get('facts')
    mock_mode = request.data.get('mock_mode', False)
    stop_after = request.data.get('stop_after', None)  # 'outline', 'sections', etc.

    try:
        if mock_mode:
            from services.pipeline.work_types import WORK_TYPE_REGISTRY

            preset = WORK_TYPE_REGISTRY.get(work_type)
            if not preset:
                preset = WORK_TYPE_REGISTRY['referat']

            mock_outline = {
                "chapters": [
                    {
                        "key": "intro",
                        "title": "Введение",
                        "sections": [
                            {"key": "intro.main", "title": "Введение", "target_words": 500}
                        ]
                    },
                    {
                        "key": "chapter1",
                        "title": "Теоретическая часть",
                        "sections": [
                            {"key": "chapter1.overview", "title": "Обзор предметной области", "target_words": 800},
                            {"key": "chapter1.tech", "title": "Анализ технологий", "target_words": 600}
                        ]
                    },
                    {
                        "key": "chapter2",
                        "title": "Практическая часть",
                        "sections": [
                            {"key": "chapter2.design", "title": "Проектирование системы", "target_words": 700},
                            {"key": "chapter2.impl", "title": "Реализация", "target_words": 800}
                        ]
                    },
                    {
                        "key": "conclusion",
                        "title": "Заключение",
                        "sections": [
                            {"key": "conclusion.main", "title": "Заключение", "target_words": 400}
                        ]
                    }
                ]
            }

            mock_sections = []
            total_words = 0
            for chapter in mock_outline['chapters']:
                for section in chapter['sections']:
                    words = section['target_words']
                    mock_sections.append({
                        "key": section['key'],
                        "chapter_key": chapter['key'],
                        "title": section['title'],
                        "text": f"[Содержимое секции '{section['title']}' - {words} слов]\n\n" + "Lorem ipsum... " * (words // 10),
                        "word_count": words
                    })
                    total_words += words

            mock_literature = {
                "sources": [
                    {
                        "type": "web",
                        "citation": "Django Documentation [Электронный ресурс]. — URL: https://docs.djangoproject.com/",
                        "relevance": "technology"
                    },
                    {
                        "type": "book",
                        "citation": "Мартин Р. Чистая архитектура. — СПб.: Питер, 2018. — 352 с.",
                        "relevance": "architecture"
                    },
                    {
                        "type": "book",
                        "citation": "Гамма Э. и др. Паттерны проектирования. — СПб.: Питер, 2020. — 368 с.",
                        "relevance": "methodology"
                    }
                ]
            }

            mock_quality = {
                "total_words": total_words,
                "sections_count": len(mock_sections),
                "completeness_score": 0.85,
                "issues": []
            }

            result = {
                "outline": mock_outline,
                "sections": mock_sections,
                "sections_count": len(mock_sections),
                "total_words": total_words,
                "literature": mock_literature,
                "sources_count": len(mock_literature['sources']),
                "quality_report": mock_quality,
                "work_type": work_type,
                "profile": profile,
                "mock": True
            }
        else:
            from django.db.models import Sum
            from apps.projects.models import (
                Project, AnalysisRun, Artifact, Document, DocumentProfile, Section
            )
            from apps.llm.models import LLMCall
            from services.pipeline import DocumentRunner
            from services.pipeline.kinds import ArtifactKind
            from services.pipeline.ensure import get_success_artifact, get_outline_artifact

            test_repo_url = "https://github.com/test/pipeline-test"
            project, _ = Project.objects.get_or_create(
                repo_url=test_repo_url,
            )

            analysis_run = AnalysisRun.objects.create(
                project=project,
                status=AnalysisRun.Status.SUCCESS
            )

            if not facts:
                facts = {
                    "project_name": topic_title,
                    "languages": [{"name": "Python", "percentage": 100}],
                    "frameworks": [{"name": "Django", "version": "5.0"}],
                    "architecture": {"pattern": "MVC", "evidence": []},
                    "databases": [{"type": "SQLite"}],
                    "facts": [
                        {"id": "default_1", "text": f"Проект: {topic_title}", "tags": ["project", "overview"]},
                        {"id": "default_2", "text": "Веб-приложение на Python/Django", "tags": ["technology", "web"]},
                    ]
                }

            Artifact.objects.create(
                analysis_run=analysis_run,
                kind=Artifact.Kind.FACTS,
                data=facts
            )

            doc_profile = DocumentProfile.objects.create(
                work_type=work_type,
                topic_title=topic_title,
                topic_description=topic_description or "",
                style_level=2 if work_type == 'course' else (3 if work_type == 'diploma' else 1),
            )

            document = Document.objects.create(
                analysis_run=analysis_run,
                profile=doc_profile,
                type=work_type,
                params={"title": topic_title, "description": topic_description},
            )

            llm_count_before = LLMCall.objects.count()

            from services.pipeline.steps import ensure_outline_v2

            partial_stops = ('outline', 'intro', 'intro_theory', 'theory', 'practice', 'intro_practice')
            if stop_after in partial_stops:
                outline_artifact = ensure_outline_v2(
                    document_id=document.id,
                    force=False,
                    profile=profile,
                    mock_mode=False,
                )

                document.refresh_from_db()
                sections_created = list(document.sections.order_by('order').values('key', 'title', 'chapter_key', 'depth', 'order'))

                outline_data = outline_artifact.data_json if outline_artifact else {}

                result = {
                    "outline": outline_data,
                    "sections_created": sections_created,
                    "sections_count": len(sections_created),
                    "work_type": work_type,
                    "profile": profile,
                    "mock": False,
                    "document_id": str(document.id),
                    "stop_after": stop_after,
                }

                sections_list = []

                generate_intro = stop_after in ('intro', 'intro_theory', 'intro_practice')
                generate_theory = stop_after in ('intro_theory', 'theory')
                generate_practice = stop_after in ('practice', 'intro_practice')

                if generate_intro:
                    from services.pipeline.steps import ensure_intro_academic

                    intro_section = document.sections.filter(key='intro').first()
                    if intro_section:
                        section_artifact = ensure_intro_academic(
                            document_id=document.id,
                            force=False,
                            profile=profile,
                            mock_mode=False,
                        )

                        intro_section.refresh_from_db()
                        intro_text = intro_section.text_current or ""
                        result["intro_section"] = {
                            "key": intro_section.key,
                            "title": intro_section.title,
                            "text": intro_text,
                            "word_count": len(intro_text.split()),
                            "validation": section_artifact.meta.get("validation", {}),
                        }
                        sections_list.append({
                            "key": intro_section.key,
                            "chapter_key": intro_section.chapter_key or "intro",
                            "title": intro_section.title,
                            "text": intro_text,
                            "word_count": len(intro_text.split()),
                        })

                if generate_theory:
                    from services.pipeline.steps import ensure_theory_section

                    theory_sections = document.sections.filter(chapter_key='theory').order_by('order')
                    theory_results = []

                    parent_keys_with_children = set(
                        document.sections.filter(
                            chapter_key='theory',
                            parent_key__isnull=False
                        ).exclude(parent_key='').values_list('parent_key', flat=True)
                    )

                    for section in theory_sections:
                        if section.key in parent_keys_with_children:
                            theory_results.append({
                                "key": section.key,
                                "chapter_key": section.chapter_key,
                                "title": section.title,
                                "text": "",
                                "word_count": 0,
                                "is_parent": True,
                            })
                            sections_list.append({
                                "key": section.key,
                                "chapter_key": section.chapter_key,
                                "title": section.title,
                                "text": "",
                                "word_count": 0,
                                "is_parent": True,
                            })
                            continue

                        section_artifact = ensure_theory_section(
                            document_id=document.id,
                            section_key=section.key,
                            force=False,
                            profile=profile,
                            mock_mode=False,
                        )
                        section.refresh_from_db()
                        section_text = section.text_current or ""
                        theory_results.append({
                            "key": section.key,
                            "chapter_key": section.chapter_key,
                            "title": section.title,
                            "text": section_text,
                            "word_count": len(section_text.split()),
                            "validation": section_artifact.meta.get("validation", {}),
                            "depth": section.depth,
                            "parent_key": section.parent_key,
                        })
                        sections_list.append({
                            "key": section.key,
                            "chapter_key": section.chapter_key,
                            "title": section.title,
                            "text": section_text,
                            "word_count": len(section_text.split()),
                            "depth": section.depth,
                            "parent_key": section.parent_key,
                        })

                    result["theory_sections"] = theory_results

                if generate_practice:
                    from services.pipeline.steps import ensure_practice_section

                    practice_sections = document.sections.filter(chapter_key='practice').order_by('order')
                    practice_results = []

                    parent_keys_with_children = set(
                        document.sections.filter(
                            chapter_key='practice',
                            parent_key__isnull=False
                        ).exclude(parent_key='').values_list('parent_key', flat=True)
                    )

                    for section in practice_sections:
                        if section.key in parent_keys_with_children:
                            practice_results.append({
                                "key": section.key,
                                "chapter_key": section.chapter_key,
                                "title": section.title,
                                "text": "",
                                "word_count": 0,
                                "is_parent": True,
                            })
                            sections_list.append({
                                "key": section.key,
                                "chapter_key": section.chapter_key,
                                "title": section.title,
                                "text": "",
                                "word_count": 0,
                                "is_parent": True,
                            })
                            continue

                        section_artifact = ensure_practice_section(
                            document_id=document.id,
                            section_key=section.key,
                            force=False,
                            profile=profile,
                            mock_mode=False,
                        )
                        section.refresh_from_db()
                        section_text = section.text_current or ""
                        practice_results.append({
                            "key": section.key,
                            "chapter_key": section.chapter_key,
                            "title": section.title,
                            "text": section_text,
                            "word_count": len(section_text.split()),
                            "validation": section_artifact.meta.get("validation", {}),
                            "depth": section.depth,
                            "parent_key": section.parent_key,
                        })
                        sections_list.append({
                            "key": section.key,
                            "chapter_key": section.chapter_key,
                            "title": section.title,
                            "text": section_text,
                            "word_count": len(section_text.split()),
                            "depth": section.depth,
                            "parent_key": section.parent_key,
                        })

                    result["practice_sections"] = practice_results

                result["sections"] = sections_list

                new_llm_calls = LLMCall.objects.filter(created_at__gte=document.created_at)
                total_tokens = sum(c.meta.get('total_tokens', 0) for c in new_llm_calls if c.meta)
                total_cost = sum(c.meta.get('cost_estimate', 0) for c in new_llm_calls if c.meta)
                models_used = list(set(c.model for c in new_llm_calls if c.model))

                result["usage"] = {
                    "llm_calls": new_llm_calls.count(),
                    "total_tokens": total_tokens,
                    "total_cost": round(total_cost, 4),
                    "model": models_used[0] if len(models_used) == 1 else ", ".join(models_used) if models_used else "unknown",
                }

                return Response({
                    "status": "success",
                    "result": result,
                    "error": None
                })

            runner = DocumentRunner(
                document_id=document.id,
                profile=profile,
                mock_mode=False,
            )
            run_result = runner.run_full()

            new_llm_calls = LLMCall.objects.filter(
                created_at__gte=document.created_at
            )
            total_tokens = 0
            total_cost = 0.0
            llm_calls_count = new_llm_calls.count()
            models_used = set()
            for call in new_llm_calls:
                if call.model:
                    models_used.add(call.model)
                if call.meta:
                    total_tokens += call.meta.get('total_tokens', 0)
                    total_cost += call.meta.get('cost_estimate', 0)
            models_used = list(models_used)

            outline_artifact = get_outline_artifact(document.id)
            outline_data = outline_artifact.data_json if outline_artifact else {}

            sections_data = []
            total_words = 0
            for section in Section.objects.filter(document=document).order_by('order'):
                text = section.edited_text or section.enriched_text or section.text_current or ""
                word_count = len(text.split())
                sections_data.append({
                    "key": section.key,
                    "chapter_key": section.chapter_key,
                    "title": section.title,
                    "text": text,
                    "word_count": word_count
                })
                total_words += word_count

            literature_artifact = get_success_artifact(document.id, ArtifactKind.LITERATURE.value)
            literature_data = literature_artifact.data_json if literature_artifact else {}

            quality_artifact = get_success_artifact(document.id, ArtifactKind.QUALITY_REPORT.value)
            quality_data = quality_artifact.data_json if quality_artifact else {}

            result = {
                "outline": outline_data,
                "sections": sections_data,
                "sections_count": len(sections_data),
                "total_words": total_words,
                "literature": literature_data,
                "sources_count": len(literature_data.get('sources', [])),
                "quality_report": quality_data,
                "work_type": work_type,
                "profile": profile,
                "mock": False,
                "document_id": str(document.id),
                "run_result": {
                    "success": run_result.success,
                    "artifacts_created": run_result.artifacts_created,
                    "artifacts_cached": run_result.artifacts_cached,
                    "errors": run_result.errors,
                    "duration_ms": run_result.duration_ms,
                },
                "usage": {
                    "llm_calls": llm_calls_count,
                    "total_tokens": total_tokens,
                    "total_cost": round(total_cost, 4),
                    "model": models_used[0] if len(models_used) == 1 else ", ".join(models_used) if models_used else "unknown",
                }
            }

        return Response({
            "status": "success",
            "result": result,
            "error": None
        })

    except Exception as e:
        return Response(
            {"status": "error", "error": str(e), "result": None},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
def generate_docx_api(request):
    from io import BytesIO
    from django.http import HttpResponse
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data = request.data
    if not data:
        return Response({"error": "No data provided"}, status=400)

    outline = data.get('outline', {})
    sections = data.get('sections', [])
    title = outline.get('title', 'Документ')
    chapters = outline.get('chapters', [])

    sections_by_key = {}
    for s in sections:
        key = s.get('key', s.get('chapter_key', ''))
        sections_by_key[key] = s

    sections_by_chapter_key = {}
    for s in sections:
        chapter_key = s.get('chapter_key', '')
        if chapter_key not in sections_by_chapter_key:
            sections_by_chapter_key[chapter_key] = []
        sections_by_chapter_key[chapter_key].append(s)

    doc = DocxDocument()

    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(14)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.first_line_indent = Cm(1.25)

    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Times New Roman'
    style_h1.font.size = Pt(14)
    style_h1.font.bold = True
    style_h1.paragraph_format.space_before = Pt(24)
    style_h1.paragraph_format.space_after = Pt(0)
    style_h1.paragraph_format.first_line_indent = Cm(0)

    style_h2 = doc.styles['Heading 2']
    style_h2.font.name = 'Times New Roman'
    style_h2.font.size = Pt(14)
    style_h2.font.bold = True
    style_h2.paragraph_format.space_before = Pt(12)
    style_h2.paragraph_format.space_after = Pt(0)
    style_h2.paragraph_format.first_line_indent = Cm(0)

    style_h3 = doc.styles['Heading 3']
    style_h3.font.name = 'Times New Roman'
    style_h3.font.size = Pt(14)
    style_h3.font.bold = True
    style_h3.paragraph_format.space_before = Pt(8)
    style_h3.paragraph_format.space_after = Pt(0)
    style_h3.paragraph_format.first_line_indent = Cm(0)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    h1 = doc.add_heading(title, level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    toc_para = doc.add_paragraph('Содержание')
    toc_para.paragraph_format.first_line_indent = Cm(0)
    toc_para.runs[0].italic = True
    toc_para.paragraph_format.space_before = Pt(12)

    for chapter in chapters:
        if chapter.get('is_auto') and chapter.get('key') == 'toc':
            continue
        chapter_title = chapter.get('title', '')
        p = doc.add_paragraph(chapter_title)
        p.paragraph_format.first_line_indent = Cm(0)

        for sec in chapter.get('sections', []):
            sec_title = sec.get('title', '')
            if sec_title:
                p = doc.add_paragraph('    ' + sec_title)
                p.paragraph_format.first_line_indent = Cm(0)

            for subsec in sec.get('subsections', []):
                subsec_title = subsec.get('title', '')
                if subsec_title:
                    p = doc.add_paragraph('        ' + subsec_title)
                    p.paragraph_format.first_line_indent = Cm(0)

    doc.add_page_break()

    for chapter in chapters:
        chapter_key = chapter.get('key', '')
        chapter_title = chapter.get('title', '')

        if chapter.get('is_auto'):
            continue

        chapter_sections = chapter.get('sections', [])

        if not chapter_sections:
            doc.add_heading(chapter_title, level=1)
            sections_in_chapter = sections_by_chapter_key.get(chapter_key, [])
            for sec_data in sections_in_chapter:
                _add_section_text_simple(doc, sec_data.get('text', ''))
        else:
            doc.add_heading(chapter_title, level=1)

            for sec in chapter_sections:
                sec_key = sec.get('key', '')
                sec_title = sec.get('title', '')
                sec_subsections = sec.get('subsections', [])

                if sec_subsections:
                    doc.add_heading(sec_title, level=2)

                    for subsec in sec_subsections:
                        subsec_key = subsec.get('key', '')
                        subsec_title = subsec.get('title', '')

                        doc.add_heading(subsec_title, level=3)

                        subsec_data = sections_by_key.get(subsec_key)
                        if subsec_data:
                            _add_section_text_simple(doc, subsec_data.get('text', ''))
                        else:
                            sections_in_chapter = sections_by_chapter_key.get(chapter_key, [])
                            for cs in sections_in_chapter:
                                if cs.get('key') == subsec_key or cs.get('title') == subsec_title:
                                    _add_section_text_simple(doc, cs.get('text', ''))
                                    break
                else:
                    doc.add_heading(sec_title, level=2)

                    sec_data = sections_by_key.get(sec_key)
                    if sec_data:
                        _add_section_text_simple(doc, sec_data.get('text', ''))
                    else:
                        sections_in_chapter = sections_by_chapter_key.get(chapter_key, [])
                        for cs in sections_in_chapter:
                            if cs.get('key') == sec_key or cs.get('title') == sec_title:
                                _add_section_text_simple(doc, cs.get('text', ''))
                                break

    if not chapters:
        for sec_data in sections:
            sec_title = sec_data.get('title', '')
            doc.add_heading(sec_title, level=2)
            _add_section_text_simple(doc, sec_data.get('text', ''))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="document.docx"'

    return response


def _add_section_text_simple(doc, text: str):
    if not text:
        return

    text = text.replace('**', '').replace('*', '').replace('`', '')
    text = text.replace('# ', '').replace('## ', '').replace('### ', '')

    paragraphs = text.split('\n\n')
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        lines = para_text.split('\n')
        combined = ' '.join(line.strip() for line in lines if line.strip())

        if combined:
            doc.add_paragraph(combined)
